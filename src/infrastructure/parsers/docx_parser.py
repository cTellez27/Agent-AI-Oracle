import uuid
from pathlib import Path
from typing import Union, BinaryIO
from docx import Document as DocxDocument

from src.domain.entities.document import Document, DocumentMetadata
from src.domain.interfaces.parser_interface import BaseParserInterface


class DocxParser(BaseParserInterface):
    """Extractor de texto para archivos Word (.docx)."""

    def supports_extension(self, extension: str) -> bool:
        return extension.lower() in [".docx", ".doc"]

    def parse(self, file_source: Union[str, Path, BinaryIO], file_name: str, category: str = "General") -> Document:
        try:
            doc = DocxDocument(file_source)
            paragraphs = []

            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    paragraphs.append(p.text.strip())

            # Extraer también texto de tablas en el documento
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        paragraphs.append(" | ".join(row_data))

            full_content = "\n\n".join(paragraphs) if paragraphs else "Documento Word sin texto relevante."

            size_bytes = 0
            if isinstance(file_source, (str, Path)):
                size_bytes = Path(file_source).stat().st_size
            elif hasattr(file_source, "seek") and hasattr(file_source, "tell"):
                file_source.seek(0, 2)
                size_bytes = file_source.tell()
                file_source.seek(0)

            metadata = DocumentMetadata(
                file_name=file_name,
                file_type="DOCX",
                file_size_bytes=size_bytes,
                category=category,
                additional_info={"num_paragraphs": len(doc.paragraphs), "num_tables": len(doc.tables)}
            )

            return Document(
                id=str(uuid.uuid4()),
                content=full_content,
                metadata=metadata
            )
        except Exception as e:
            raise ValueError(f"Error al procesar el archivo Word '{file_name}': {str(e)}")
